from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.2)
section.bottom_margin = Cm(2.2)

# ── Background: black page ───────────────────────────────────────────────────
def set_page_bg(doc, hex_color="000000"):
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), hex_color)
    doc.element.insert(0, bg)
    settings = doc.settings.element
    ds = OxmlElement('w:displayBackgroundShape')
    settings.insert(0, ds)

set_page_bg(doc, "000000")

# ── Color palette ────────────────────────────────────────────────────────────
BLACK   = RGBColor(0x00, 0x00, 0x00)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
EMERALD = RGBColor(0x10, 0xB9, 0x81)
GRAY    = RGBColor(0x6B, 0x72, 0x80)
LGRAY   = RGBColor(0xD1, 0xD5, 0xDB)
DGRAY   = RGBColor(0x11, 0x11, 0x11)

# ── Helpers ──────────────────────────────────────────────────────────────────
def clear_default_styles(doc):
    style = doc.styles['Normal']
    style.font.color.rgb = WHITE
    style.font.name = 'Courier New'
    style.font.size = Pt(10)

def shade_paragraph(para, hex_fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

def add_border_bottom(para, color="1F2937", size="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def para_space(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'),  str(after))
    pPr.append(sp)

clear_default_styles(doc)

# ── HEADER ───────────────────────────────────────────────────────────────────
header_p = doc.add_paragraph()
shade_paragraph(header_p, "000000")
para_space(header_p, before=0, after=80)
header_p.paragraph_format.left_indent = Cm(0)

run = header_p.add_run("DOJI")
run.font.name  = 'Courier New'
run.font.size  = Pt(22)
run.font.bold  = True
run.font.color.rgb = WHITE

run2 = header_p.add_run(" FUNDING")
run2.font.name  = 'Courier New'
run2.font.size  = Pt(22)
run2.font.bold  = False
run2.font.color.rgb = EMERALD

# ── Thin green rule ──────────────────────────────────────────────────────────
rule = doc.add_paragraph()
shade_paragraph(rule, "000000")
para_space(rule, before=0, after=120)
add_border_bottom(rule, color="10B981", size="4")

# ── Document label ───────────────────────────────────────────────────────────
label_p = doc.add_paragraph()
shade_paragraph(label_p, "000000")
para_space(label_p, before=0, after=40)
lr = label_p.add_run("01 | STRATEGY")
lr.font.name  = 'Courier New'
lr.font.size  = Pt(8)
lr.font.color.rgb = EMERALD

# ── Title ────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
shade_paragraph(title_p, "000000")
para_space(title_p, before=0, after=60)
tr = title_p.add_run("Cold-Start Procedure")
tr.font.name  = 'Courier New'
tr.font.size  = Pt(28)
tr.font.bold  = True
tr.font.color.rgb = WHITE

sub_p = doc.add_paragraph()
shade_paragraph(sub_p, "000000")
para_space(sub_p, before=0, after=200)
sr = sub_p.add_run("From go-live to growth — DXTrade launch playbook")
sr.font.name  = 'Courier New'
sr.font.size  = Pt(11)
sr.font.color.rgb = GRAY

# ── Section helper ───────────────────────────────────────────────────────────
def add_phase(doc, num, title, goal, bullets):
    # Phase header
    ph = doc.add_paragraph()
    shade_paragraph(ph, "000000")
    para_space(ph, before=200, after=30)
    add_border_bottom(ph, color="1F2937", size="4")

    n_run = ph.add_run(f"{num}  ")
    n_run.font.name  = 'Courier New'
    n_run.font.size  = Pt(8)
    n_run.font.color.rgb = EMERALD
    n_run.font.bold  = True

    t_run = ph.add_run(title.upper())
    t_run.font.name  = 'Courier New'
    t_run.font.size  = Pt(14)
    t_run.font.bold  = True
    t_run.font.color.rgb = WHITE

    # Timeline chip
    tp = doc.add_paragraph()
    shade_paragraph(tp, "111111")
    para_space(tp, before=30, after=40)
    tp.paragraph_format.left_indent = Cm(0)
    chip = tp.add_run(f"  GOAL — {goal}  ")
    chip.font.name  = 'Courier New'
    chip.font.size  = Pt(8)
    chip.font.color.rgb = EMERALD

    # Bullets
    for b in bullets:
        bp = doc.add_paragraph()
        shade_paragraph(bp, "000000")
        para_space(bp, before=20, after=20)
        bp.paragraph_format.left_indent = Cm(0.5)

        dot = bp.add_run("▸  ")
        dot.font.name  = 'Courier New'
        dot.font.size  = Pt(9)
        dot.font.color.rgb = EMERALD

        txt = bp.add_run(b)
        txt.font.name  = 'Courier New'
        txt.font.size  = Pt(9)
        txt.font.color.rgb = LGRAY

# ── Phases ───────────────────────────────────────────────────────────────────
add_phase(doc, "01", "Soft Launch", "Validate the full flow end-to-end before any traffic", [
    "Create 3–5 internal test accounts and run a full challenge cycle (purchase → evaluation → funded → payout)",
    "Stress-test DXTrade rules: drawdown limits, profit targets, account resets",
    "Verify webhook/API connections between your dashboard and DXTrade",
    "QA the payment flow (deposit → account creation → confirmation email)",
    "Fix all blockers before any public exposure",
])

add_phase(doc, "02", "Closed Beta", "Real users, controlled volume", [
    "Invite 20–50 traders (friends, Discord members, early signups)",
    "Offer a discounted or free challenge in exchange for feedback",
    "Monitor support tickets closely — identify friction points",
    "Validate payout flow with at least 1–2 real payouts processed",
    "Collect testimonials from satisfied beta traders",
])

add_phase(doc, "03", "Public Launch", "Controlled growth with measurable CAC", [
    "Announce on X/Twitter, Discord, prop trading communities (Reddit, Telegram)",
    "Run a launch promo — limited-time discount on first challenge",
    "Set up affiliate tracking if your affiliate program is ready",
    "Monitor DXTrade server load and dashboard performance under real traffic",
    "Track key metrics from day 1: conversion rate, challenge pass rate, churn",
])

add_phase(doc, "04", "Growth Loop", "Sustained, compounding acquisition", [
    "Affiliate program → prop trading influencers drive volume",
    "Leaderboard → public rankings build social proof and retention",
    "Retargeting → users who started checkout but didn't convert",
    "Discord community → reduce churn, build brand loyalty",
])

# ── Footer rule ──────────────────────────────────────────────────────────────
fr = doc.add_paragraph()
shade_paragraph(fr, "000000")
para_space(fr, before=200, after=40)
add_border_bottom(fr, color="1F2937", size="4")

fp = doc.add_paragraph()
shade_paragraph(fp, "000000")
para_space(fp, before=0, after=0)
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
f1 = fp.add_run("dojifunding.com")
f1.font.name  = 'Courier New'
f1.font.size  = Pt(8)
f1.font.color.rgb = GRAY

f2 = fp.add_run("   —   CONFIDENTIAL")
f2.font.name  = 'Courier New'
f2.font.size  = Pt(8)
f2.font.color.rgb = EMERALD

# ── Save ─────────────────────────────────────────────────────────────────────
out = r"c:\Users\Phidias Research\Desktop\Doji funding\Doji_ColdStart_Procedure.docx"
doc.save(out)
print("Saved:", out)
