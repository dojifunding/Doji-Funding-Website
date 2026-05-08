#!/usr/bin/env python3
"""Génère les deux documents Word (.docx) pour la documentation technique Doji Funding."""

import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
GREEN  = RGBColor(0x10, 0xB9, 0x81)
BLACK  = RGBColor(0x00, 0x00, 0x00)
DARK   = RGBColor(0x1A, 0x1A, 0x1A)
GRAY   = RGBColor(0x44, 0x44, 0x44)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG= '111111'
CODE_FG= RGBColor(0xE8, 0xE8, 0xE8)

# ─── XML helpers ──────────────────────────────────────────────────────────────

def shade_para(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def left_border(para, color='10B981', sz='28'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    e = OxmlElement('w:left')
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
    e.set(qn('w:space'), '6'); e.set(qn('w:color'), color)
    pBdr.append(e); pPr.append(pBdr)

def bottom_border(para, color='10B981', sz='12'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    e = OxmlElement('w:bottom')
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
    e.set(qn('w:space'), '4'); e.set(qn('w:color'), color)
    pBdr.append(e); pPr.append(pBdr)

def cell_shade(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

# ─── Content helpers ──────────────────────────────────────────────────────────

def strip_html(s):
    s = re.sub(r'<[^>]+>', '', s)
    for e, r in [('&lt;','<'),('&gt;','>'),('&amp;','&'),('&nbsp;',' '),
                 ('&#39;',"'"),('&quot;','"'),('&lt;!--','<!--'),('--&gt;','-->')]:
        s = s.replace(e, r)
    return s

def h_cover(doc, title, subtitle, meta_pairs):
    """Cover page block."""
    doc.add_paragraph()
    doc.add_paragraph()

    # Badge
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Documentation Technique')
    r.font.name = 'Calibri'; r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = BLACK; r.font.all_caps = True
    shade_para(p, '10B981')
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = 'Calibri'; r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)
    bottom_border(p)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.name = 'Calibri'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(20)

    # Meta
    for label, value in meta_pairs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + ' '); r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = RGBColor(0x33,0x33,0x33)
        r2 = p.add_run(value); r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x55,0x55,0x55)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

def h2(doc, num, title):
    p = doc.add_heading(level=2); p.clear()
    r0 = p.add_run(f' {num} ')
    r0.font.name = 'Calibri'; r0.font.bold = True; r0.font.size = Pt(11)
    r0.font.color.rgb = BLACK
    shade_para(p, '10B981')  # green badge background via shading on whole para — use inline XML instead
    # Reset shade, use inline approach for badge
    # Redo: just prefix number in brackets, green color
    p.clear()
    r0 = p.add_run(f'[{num}] ')
    r0.font.name = 'Calibri'; r0.font.bold = True; r0.font.size = Pt(14)
    r0.font.color.rgb = GREEN
    r1 = p.add_run(title)
    r1.font.name = 'Calibri'; r1.font.bold = True; r1.font.size = Pt(16)
    r1.font.color.rgb = BLACK
    bottom_border(p)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)

def h3(doc, title):
    p = doc.add_heading(level=3); p.clear()
    r0 = p.add_run('▸ ')
    r0.font.name = 'Calibri'; r0.font.size = Pt(12); r0.font.color.rgb = GREEN
    r1 = p.add_run(title)
    r1.font.name = 'Calibri'; r1.font.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)

def h4(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(title.upper())
    r.font.name = 'Calibri'; r.font.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x44,0x44,0x44)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)

def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(strip_html(text))
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(8)
    return p

def code(doc, text, border_color='10B981'):
    lines = strip_html(text).strip('\n').split('\n')
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        r = p.add_run(line if line.strip() else ' ')
        r.font.name = 'Courier New'; r.font.size = Pt(9); r.font.color.rgb = CODE_FG
        shade_para(p, CODE_BG)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_before = Pt(1 if i > 0 else 8)
        p.paragraph_format.space_after = Pt(1 if i < len(lines)-1 else 12)
        if i == 0:
            left_border(p, border_color)

def box(doc, text, type='info'):
    palettes = {
        'info': ('DCF5EC', '10B981'),
        'warn': ('FFF8E6', 'F0A500'),
        'tip' : ('EEF4FF', '4A90D9'),
        'key' : ('FFF0F0', 'E74C3C'),
    }
    bg, bd = palettes.get(type, ('DCF5EC','10B981'))
    p = doc.add_paragraph()
    r = p.add_run(strip_html(text))
    r.font.name = 'Calibri'; r.font.size = Pt(10.5); r.font.color.rgb = DARK
    shade_para(p, bg); left_border(p, bd)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(12)

def step(doc, n, title, desc=''):
    p = doc.add_paragraph()
    r0 = p.add_run(f'  {n}  ')
    r0.font.name = 'Courier New'; r0.font.bold = True; r0.font.size = Pt(10)
    r0.font.color.rgb = BLACK
    shade_para(p, '10B981')
    # Re-do without shading hack:
    p.clear()
    r0 = p.add_run(f'[{n}] ')
    r0.font.bold = True; r0.font.size = Pt(11); r0.font.color.rgb = GREEN
    r1 = p.add_run(title)
    r1.font.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    if desc:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(strip_html(desc))
        r2.font.name = 'Calibri'; r2.font.size = Pt(10.5); r2.font.color.rgb = GRAY
        p2.paragraph_format.left_indent = Inches(0.4)
        p2.paragraph_format.space_after = Pt(8)

def table_2col(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header row
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        cell_shade(c, '10B981')
        run = c.paragraphs[0].runs[0]
        run.font.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = BLACK; run.font.all_caps = True
    # Data rows
    for ri, row in enumerate(rows):
        dr = t.rows[ri+1]
        bg = 'F9F9F9' if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = dr.cells[ci]; c.text = val
            cell_shade(c, bg)
            run = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
    t.autofit = True
    # Spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def divider(doc):
    p = doc.add_paragraph()
    bottom_border(p, 'EEEEEE', '6')
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(16)

def footer(doc):
    divider(doc)
    p = doc.add_paragraph()
    r = p.add_run('Doji Funding® — Documentation Interne Technique\nGénéré par Claude Code · Ne pas distribuer hors de l\'équipe')
    r.font.name = 'Calibri'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 : HERO CANVAS
# ══════════════════════════════════════════════════════════════════════════════

def create_hero_canvas_doc():
    doc = Document()

    # Page margins
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3);  s.right_margin = Cm(3)

    # ── Cover ──
    h_cover(doc,
        'Hero Canvas\nAnimations 3D',
        'Globe · Cube · Pyramide · Diamant — Dot Grid interactif\nDoji Funding® — Frontend',
        [
            ('Fichier principal :', 'assets/js/hero-canvas.js'),
            ('Dépendances :', 'Aucune — Canvas 2D natif (pas WebGL)'),
            ('API utilisée :', 'Canvas 2D'),
        ]
    )

    # ── S1 ──
    h2(doc, '01', 'Vue d\'ensemble — Architecture')
    body(doc, 'Tout le code réside dans un seul fichier auto-exécutable (IIFE) : assets/js/hero-canvas.js. '
         'Il n\'utilise aucune bibliothèque externe — pas de Three.js, pas de GSAP, rien d\'autre que l\'API Canvas 2D native du navigateur.')
    box(doc, 'Principe fondamental : La "3D" est simulée avec de simples calculs de trigonométrie et une projection perspective à la main. '
        'Le navigateur n\'utilise que ctx.arc() pour dessiner des cercles — chaque point 3D devient un petit disque 2D sur le canvas.', 'info')
    body(doc, 'Chaque frame (60 fois par seconde via requestAnimationFrame), le moteur effectue deux passes :')
    code(doc, '''\
FRAME N  (appelé ~60x/sec par requestAnimationFrame)

  PASSE 1 — DOT GRID
  Pour chaque point de la grille :
    1. Calcul du warp (anneau autour de la shape)
    2. Calcul répulsion souris
    3. Application spring + friction
    4. Dessin du point (blanc ou émeraude)

  PASSE 2 — SHAPE 3D
  Pour chaque point 3D de la shape :
    1. Rotation Y puis X
    2. Projection perspective (division par Z)
    3. Tri par profondeur (painter\'s algorithm)
    4. Dessin du point (couleur + taille selon Z)''')
    body(doc, 'L\'objet principal HeroCanvas gère un canvas. S\'il y a plusieurs canvas sur la page, autant d\'instances sont créées.')

    # ── S2 ──
    h2(doc, '02', 'Le Dot Grid — Grille de points interactifs')
    body(doc, 'La grille couvre tout le canvas. Les points sont espacés de 22px et ont un rayon de 1.5px. '
         'Chaque point est une instance de l\'objet Dot.')
    code(doc, '''\
var GRID  = 22;    // espacement entre les points en pixels
var DOT_R = 1.5;   // rayon de chaque point en pixels

function Dot(ox, oy) {
  this.ox    = ox; this.oy    = oy;  // position d\'origine (fixe)
  this.x     = ox; this.y     = oy;  // position courante (déplacée)
  this.vx    = 0;  this.vy    = 0;   // vélocité
  this.green = 0;                    // 0 = blanc, 1 = émeraude
}''')

    h3(doc, 'Physique spring (ressort)')
    body(doc, 'Chaque point est attiré vers sa position d\'origine par un ressort. '
         'La force est proportionnelle à la distance, la friction dissipe l\'énergie à chaque frame.')
    code(doc, '''\
// La "cible" peut être décalée par le warp (voir plus bas)
this.vx += (target_x - this.x) * SPRING;   // SPRING = 0.07
this.vy += (target_y - this.y) * SPRING;

this.vx *= FRIC;   // FRIC = 0.82 → amortissement (freine la vélocité)
this.vy *= FRIC;

this.x  += this.vx;   // déplace le point
this.y  += this.vy;''')

    h3(doc, 'Interaction souris — scatter émeraude')
    body(doc, 'Quand le curseur entre dans un rayon de 115px autour d\'un point, celui-ci est repoussé et devient émeraude. '
         'La couleur s\'estompe progressivement après que le curseur s\'éloigne.')
    code(doc, '''\
var M_RAD   = 115;   // rayon d\'influence de la souris (px)
var M_FORCE = 10;    // force de répulsion

if (mouseOn) {
  var dx = this.x - mouseX;
  var dy = this.y - mouseY;
  var d2 = dx*dx + dy*dy;

  if (d2 < M_RAD * M_RAD) {
    var d = Math.sqrt(d2) || 1;
    var f = (M_RAD - d) / M_RAD;   // force décroît avec la distance
    this.vx   += (dx / d) * f * M_FORCE;
    this.vy   += (dy / d) * f * M_FORCE;
    this.green = 1;   // devient émeraude instantanément
  }
}

// Fade-out de la couleur émeraude
if (this.green > 0) {
  this.green *= 0.96;   // diminue de 4% par frame → ~70 frames pour disparaître
  if (this.green < 0.02) this.green = 0;
}''')

    h3(doc, 'Warp spacetime — anneau de distorsion')
    body(doc, 'C\'est l\'effet le plus distinctif : les points se déforment en anneau autour du centre de la shape. '
         'C\'est une fonction gaussienne — les points trop proches et trop loin ne bougent presque pas, '
         'seul l\'anneau intermédiaire est poussé vers l\'extérieur.')
    code(doc, '''\
var SP_PUSH = 58;    // amplitude max du warp (px)
var SP_RAD  = 148;   // rayon de référence (centre de l\'anneau)

// sdx/sdy = vecteur du point vers le centre de la shape
var n   = distance_au_centre / SP_RAD;
var amp = SP_PUSH * n * Math.exp(-n * n * 0.7);

// La cible est déplacée dans la direction radiale d\'une amplitude "amp"
target_x = origine_x + (direction_x / distance) * amp;
target_y = origine_y + (direction_y / distance) * amp;''')
    box(doc, 'Astuce : La formule n × e^(−n²×0.7) produit un pic à n ≈ 0.85, soit environ 85% du rayon SP_RAD. '
        'C\'est ce qui crée l\'anneau visible. En modifiant 0.7, on change la largeur du pic.', 'tip')

    # ── S3 ──
    h2(doc, '03', 'Les Shapes 3D — Principe commun')

    h3(doc, 'Représentation des points')
    body(doc, 'Chaque shape est simplement un tableau JavaScript de points 3D. '
         'Un point est un tableau de 3 valeurs [X, Y, Z] normalisées dans l\'espace −1 à +1.')
    code(doc, '''\
// Exemple : 4 points d\'un carré dans le plan Z=0
var points = [
  [-1, -1,  0],
  [ 1, -1,  0],
  [ 1,  1,  0],
  [-1,  1,  0]
];''')

    h3(doc, 'Interpolation des arêtes — edgePts(A, B, n)')
    body(doc, 'Les shapes ne sont pas des meshes remplis — ce sont des wireframes faits de points. '
         'La fonction edgePts crée n points régulièrement espacés entre deux sommets A et B, simulant ainsi une arête visible.')
    code(doc, '''\
function edgePts(pa, pb, n) {
  var pts = [];
  for (var i = 0; i <= n; i++) {
    var t = i / n;   // t varie de 0 à 1
    pts.push([
      pa[0] + (pb[0] - pa[0]) * t,   // X interpolé
      pa[1] + (pb[1] - pa[1]) * t,   // Y interpolé
      pa[2] + (pb[2] - pa[2]) * t    // Z interpolé
    ]);
  }
  return pts;
}

// Exemple : arête du cube avec 44 points
pts = pts.concat(edgePts([-1,-1,-1], [1,-1,-1], 44));''')

    h3(doc, 'Projection perspective')
    body(doc, 'La projection transforme un point 3D en coordonnées 2D sur le canvas. '
         'Plus un point est loin (Z négatif = fond de la scène), plus il apparaît petit et sombre. '
         'Plus il est proche (Z positif), plus il est grand et lumineux.')
    code(doc, '''\
var sc  = 120;         // scale de base (rayon de la shape en pixels)
var fov = sc * 3;      // champ de vision (distance oeil → écran)

// 1. Rotation (détaillée section suivante)
var q = rY(point, rotY);
    q = rX(q,     rotX);

// 2. Translation Z pour éviter division par zéro
var qz = q[2] * sc + sc * 2.5;

// 3. Facteur de perspective : objets lointains semblent plus petits
var s = fov / qz;

// 4. Coordonnées 2D finales
var px = centreX + q[0] * sc * s;
var py = centreY + q[1] * sc * s;

// 5. Opacité et taille selon la profondeur (painter\'s shading)
var alpha  = 0.28 + (q[2] + 1) * 0.36;   // 0.28 en fond → 1.0 en avant
var radius = 1.1  + (q[2] + 1) * 0.55;   // plus petit en fond, plus grand en avant''')
    box(doc, 'Painter\'s algorithm : Avant de dessiner, tous les points sont triés par Z croissant (les plus lointains d\'abord). '
        'Ainsi les points proches sont dessinés par-dessus les lointains, créant l\'illusion de profondeur sans calcul d\'occlusion.', 'info')

    h3(doc, 'Rotation interactive')
    body(doc, 'Deux types de rotation coexistent : une auto-rotation constante, et une rotation interactive pilotée par la souris. '
         'Un lerp (interpolation linéaire) lisse la transition entre les deux.')
    code(doc, '''\
// Fonctions de rotation 3D pures
function rY(p, a) {   // rotation autour de l\'axe Y
  var c = Math.cos(a), s = Math.sin(a);
  return [c*p[0] + s*p[2],  p[1],  -s*p[0] + c*p[2]];
}
function rX(p, a) {   // rotation autour de l\'axe X
  var c = Math.cos(a), s = Math.sin(a);
  return [p[0],  c*p[1] - s*p[2],  s*p[1] + c*p[2]];
}

// Chaque frame :
autoY += 0.004;   // auto-rotation constante (rad/frame)

if (mouseOver) {
  // Souris : mappage position → angle
  tgtY = (mouseX / W - 0.5) * 0.55 * 2;   // ±55° selon X
  tgtX = (mouseY / H - 0.5) * 0.55;        // ±27.5° selon Y
} else {
  tgtY = autoY;                              // reprend l\'auto-rotation
  tgtX += (0 - tgtX) * 0.03;               // revient à 0 progressivement
}

// Lerp (interpolation douce) — ROT_LERP = 0.055
rotY += (tgtY - rotY) * 0.055;
rotX += (tgtX - rotX) * 0.055;''')

    # ── S4 — LE GLOBE ──
    h2(doc, '04', 'Le Globe — Guide détaillé de reconstruction')
    body(doc, 'Le globe est composé de deux systèmes de lignes : des parallèles (lignes horizontales de latitude) '
         'et des méridiens (grands cercles verticaux).')

    h3(doc, 'Lignes de latitude (parallèles)')
    body(doc, 'Chaque parallèle est un anneau horizontal à une hauteur Y donnée. '
         'Le rayon varie selon le cosinus de la latitude — maximal à l\'équateur (cos 0° = 1), nul aux pôles (cos 90° = 0).')
    code(doc, '''\
// Parallèles tous les 15° entre −75° et +75°
// (les pôles à ±90° ne sont pas tracés — un seul point)
for (var latD = -75; latD <= 75; latD += 15) {

  var lat = latD * Math.PI / 180;   // convertir degrés → radians

  var r  = Math.cos(lat);   // rayon horizontal du cercle à cette latitude
  var yv = Math.sin(lat);   // hauteur Y (sin donne la composante verticale)

  // Nombre de points adaptatif : plus de points sur les grands cercles
  // L\'équateur aura 52 points, les cercles polaires beaucoup moins
  var N = Math.max(12, Math.round(r * 52));

  for (var i = 0; i < N; i++) {
    var a = (i / N) * Math.PI * 2;   // angle sur le cercle (0 → 360°)
    pts.push([
      r * Math.cos(a),   // X = projection horizontale
      yv,                // Y = hauteur fixe pour ce parallèle
      r * Math.sin(a)    // Z = profondeur
    ]);
  }
}''')

    table_2col(doc,
        ['Latitude', 'r = cos(lat)', 'Nb de points'],
        [
            ['0° (équateur)', '1.00', '52 points'],
            ['±15°', '0.97', '50 points'],
            ['±30°', '0.87', '45 points'],
            ['±45°', '0.71', '37 points'],
            ['±60°', '0.50', '26 points'],
            ['±75°', '0.26', '14 points'],
        ]
    )

    h3(doc, 'Méridiens (grands cercles verticaux)')
    body(doc, 'Chaque méridien est un grand cercle vertical qui passe par les deux pôles. '
         'Chaque méridien est un cercle complet (0° à 360°), donc il dessine automatiquement les deux hémisphères. '
         'On n\'a besoin que de 9 méridiens (de 0° à 160°, pas 180°) car le méridien à 0° et celui à 180° seraient identiques.')
    code(doc, '''\
// Méridiens tous les 20° de longitude
// De 0° à 160° : 9 méridiens (180° serait identique à 0° car cercle complet)
for (var lonD = 0; lonD < 180; lonD += 20) {

  var lon = lonD * Math.PI / 180;   // convertir degrés → radians
  var M   = 52;   // points par méridien

  for (var j = 0; j <= M; j++) {
    var th = (j / M) * Math.PI * 2;   // θ : 0 → 2π (cercle complet)

    // Formule d\'un grand cercle en coordonnées sphériques :
    // x = cos(θ)·cos(lon)  →  projection sur X selon la longitude
    // y = sin(θ)           →  hauteur Y (de −1 à +1, les deux pôles)
    // z = cos(θ)·sin(lon)  →  profondeur Z selon la longitude
    pts.push([
      Math.cos(th) * Math.cos(lon),
      Math.sin(th),
      Math.cos(th) * Math.sin(lon)
    ]);
  }
}''')
    box(doc, 'Attention : Contrairement aux parallèles où l\'angle θ va de 0 à 2π horizontalement, ici θ est l\'angle vertical '
        'le long du méridien. Quand θ = 0, on est sur l\'équateur face avant ; quand θ = π/2 on est au pôle nord ; '
        'quand θ = π on est sur l\'équateur face arrière ; quand θ = 3π/2 on est au pôle sud.', 'warn')

    h3(doc, 'Code complet annoté de buildSphere()')
    code(doc, '''\
function buildSphere() {
  var pts = [];   // tableau qui contiendra tous les points 3D

  // ══════════════════════════════════════════════════════
  // PARTIE 1 : Lignes de latitude (parallèles horizontaux)
  // ══════════════════════════════════════════════════════
  for (var latD = -75; latD <= 75; latD += 15) {
    var lat = latD * Math.PI / 180;
    var r   = Math.cos(lat);   // rayon du parallèle (1 à l\'équateur, 0 aux pôles)
    var yv  = Math.sin(lat);   // hauteur Y du parallèle
    var N   = Math.max(12, Math.round(r * 52));  // nb points adaptatif

    for (var i = 0; i < N; i++) {
      var a = (i / N) * Math.PI * 2;
      pts.push([r * Math.cos(a), yv, r * Math.sin(a)]);
    }
  }

  // ══════════════════════════════════════════════════════
  // PARTIE 2 : Méridiens (grands cercles verticaux)
  // ══════════════════════════════════════════════════════
  for (var lonD = 0; lonD < 180; lonD += 20) {
    var lon = lonD * Math.PI / 180;
    var M   = 52;

    for (var j = 0; j <= M; j++) {
      var th = (j / M) * Math.PI * 2;
      pts.push([
        Math.cos(th) * Math.cos(lon),
        Math.sin(th),
        Math.cos(th) * Math.sin(lon)
      ]);
    }
  }

  return pts;   // ~900 points au total
}''')

    # ── S5 ──
    h2(doc, '05', 'Les autres shapes')
    table_2col(doc,
        ['Shape', 'data-shape', 'Méthode de construction', 'Points ~'],
        [
            ['Globe',    'sphere',   'Anneaux de latitude + méridiens trigonométriques',              '~900'],
            ['Cube',     'cube',     '8 sommets définis, 12 arêtes interpolées avec edgePts()',       '~528'],
            ['Pyramide', 'pyramid',  'Apex + 4 sommets base, arêtes + diagonales interpolées',       '~440'],
            ['Diamant',  'diamond',  '3 niveaux (table octogonale + ceinture + culet), couronne et pavillon', '~780'],
        ]
    )
    h3(doc, 'Cube — code de construction')
    code(doc, '''\
// 8 sommets d\'un cube unitaire (−1 à +1 sur chaque axe)
var C = [
  [-1,-1,-1], [1,-1,-1], [1,1,-1], [-1,1,-1],  // face arrière
  [-1,-1, 1], [1,-1, 1], [1,1, 1], [-1,1, 1]   // face avant
];

// 12 arêtes (paires d\'indices dans C)
var E = [
  [0,1],[1,2],[2,3],[3,0],   // face arrière
  [4,5],[5,6],[6,7],[7,4],   // face avant
  [0,4],[1,5],[2,6],[3,7]    // arêtes latérales
];

E.forEach(function(e) {
  pts = pts.concat(edgePts(C[e[0]], C[e[1]], 44));  // 44 points par arête
});''')

    # ── S6 ──
    h2(doc, '06', 'Intégration dans une page')
    h3(doc, '1. Le canvas HTML')
    code(doc, '''\
<!-- Conteneur parent (position: relative obligatoire) -->
<section class="hero" style="position: relative; height: 600px;">

  <canvas
    class="hero-shape-canvas"
    data-shape="sphere"    <!-- sphere | cube | pyramid | diamond -->
    data-cx="0.65"          <!-- position X du centre : 0=gauche, 1=droite -->
    data-cy="0.5"           <!-- position Y du centre : 0=haut, 1=bas -->
    data-scale="1.0">       <!-- taille (1.0 = défaut 120px de rayon) -->
  </canvas>

</section>''')
    h3(doc, '2. Le CSS du canvas')
    code(doc, '''\
.hero-shape-canvas {
  position: absolute;
  top: 0; left: 0;
  width: 100%; height: 100%;
  pointer-events: none;   /* ne bloque pas les clics sur le contenu */
  opacity: .9;
}''')
    h3(doc, '3. Le script')
    code(doc, '''\
<!-- Charger le script en bas de page ou avec defer -->
<script defer src="assets/js/hero-canvas.js"></script>''')
    box(doc, 'Auto-détection : Le script détecte automatiquement tous les éléments .hero-shape-canvas présents '
        'dans la page et crée une instance pour chacun. On peut avoir plusieurs shapes différents sur la même page.', 'info')
    h3(doc, '4. Optimisation performances')
    body(doc, 'Le moteur utilise un IntersectionObserver pour suspendre l\'animation quand le canvas n\'est pas visible '
         'à l\'écran. L\'animation reprend automatiquement quand il redevient visible. '
         'Sur mobile, la shape est automatiquement centrée (data-cx est ignoré, forcé à 0.5).')

    # ── S7 ──
    h2(doc, '07', 'Paramètres de configuration')
    table_2col(doc,
        ['Variable JS', 'Valeur', 'Rôle'],
        [
            ['GRID',      '22',    'Espacement entre les points de la grille (px)'],
            ['DOT_R',     '1.5',   'Rayon des points de la grille (px)'],
            ['M_RAD',     '115',   'Rayon d\'influence de la souris (px)'],
            ['M_FORCE',   '10',    'Force de répulsion de la souris'],
            ['SPRING',    '0.07',  'Raideur du ressort (retour à l\'origine)'],
            ['FRIC',      '0.82',  'Friction (0=stop immédiat, 1=pas d\'amortissement)'],
            ['SP_PUSH',   '58',    'Amplitude max du warp spacetime (px)'],
            ['SP_RAD',    '148',   'Rayon de référence du warp (px)'],
            ['S_SCALE',   '120',   'Rayon de base de la shape (px)'],
            ['A_ROT',     '0.004', 'Vitesse d\'auto-rotation (radians par frame)'],
            ['ROT_LERP',  '0.055', 'Lissage de la rotation (valeur basse = inertie forte)'],
            ['M_ROT_MAX', '0.55',  'Angle max de rotation par la souris (radians)'],
        ]
    )
    table_2col(doc,
        ['Attribut HTML', 'Valeur par défaut', 'Rôle'],
        [
            ['data-shape', 'sphere', 'Type de shape : sphere, cube, pyramid, diamond'],
            ['data-cx',    '0.5',    'Position X du centre (0=gauche, 1=droite, 0.65=décalé droite)'],
            ['data-cy',    '0.5',    'Position Y du centre (0=haut, 1=bas)'],
            ['data-scale', '1.0',    'Multiplicateur de taille (1.2 = 20% plus grand)'],
        ]
    )

    # ── S8 — CHECKLIST ──
    h2(doc, '08', 'Checklist — Refaire le globe from scratch')
    step(doc, 1, 'Créer le fichier hero-canvas.js',
         'Wrapper IIFE : (function() { \'use strict\'; /* ... */ })();')
    step(doc, 2, 'Déclarer les constantes de configuration',
         'GRID, DOT_R, M_RAD, M_FORCE, SPRING, FRIC, SP_PUSH, SP_RAD, S_SCALE, A_ROT, ROT_LERP, M_ROT_MAX')
    step(doc, 3, 'Implémenter l\'objet Dot',
         'Constructeur avec ox/oy/x/y/vx/vy/green + méthodes tick() et draw()')
    step(doc, 4, 'Implémenter buildSphere()',
         'Boucle latitude (parallèles) + boucle longitude (méridiens) → retourner le tableau de points 3D')
    step(doc, 5, 'Implémenter les rotations rY() et rX()',
         'Deux fonctions pures de rotation 3D utilisant cos/sin')
    step(doc, 6, 'Créer l\'objet HeroCanvas',
         'Constructor → _buildShape() + _buildGrid() + _bind() + _observe() + _tick()')
    step(doc, 7, 'Implémenter _draw()',
         '1. Mise à jour rotation | 2. clearRect | 3. Tick + draw de chaque dot | 4. Projection + tri + dessin de chaque point 3D')
    step(doc, 8, 'Bootstrap',
         'document.querySelectorAll(\'.hero-shape-canvas\').forEach(el => new HeroCanvas(el))')
    step(doc, 9, 'HTML + CSS',
         'Canvas avec classe hero-shape-canvas + attributs data-shape/cx/cy/scale + CSS position absolute + pointer-events:none')
    step(doc, 10, 'Charger le script',
         'Tag <script defer> en bas de page ou dans le head')
    box(doc, 'Conseil : Pour tester rapidement, commencer avec seulement les méridiens (supprimer la boucle des parallèles) '
        'et vérifier que les cercles verticaux tournent correctement. Ajouter les parallèles ensuite. '
        'C\'est plus facile à déboguer en deux étapes.', 'tip')

    footer(doc)
    path = os.path.join(BASE, 'Hero_Canvas_Animations_3D.docx')
    doc.save(path)
    print(f'Saved: {path}')


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2 : HIW VOXEL
# ══════════════════════════════════════════════════════════════════════════════

def create_hiw_voxel_doc():
    doc = Document()

    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3);  s.right_margin = Cm(3)

    # ── Cover ──
    h_cover(doc,
        '"Your Path to\nFunded Trading"',
        'Animations Voxel Isométriques — Section How It Works\nDoji Funding® — Frontend',
        [
            ('Fichier :', 'assets/js/hiw-voxel.js'),
            ('Dépendances :', 'Aucune — Canvas 2D (fillRect)'),
            ('Projection :', 'Isométrique + rotation Y dynamique'),
        ]
    )

    # ── S1 ──
    h2(doc, '01', 'Vue d\'ensemble — Concept et rendu')
    body(doc, 'La section "Your path to funded trading" contient 4 cartes (Choose Plan, Pass Challenge, Verify Identity, Get Funded). '
         'À chaque carte est associée une animation 3D en coin supérieur droit qui s\'active au survol de la carte. '
         'Ces animations sont rendues avec du Canvas 2D pur, sans WebGL ni bibliothèque externe.')
    box(doc, 'Différence fondamentale avec hero-canvas.js : Le hero utilise une projection perspective (les objets lointains rapetissent). '
        'Ici, on utilise une projection isométrique (les objets conservent leur taille quelle que soit leur profondeur). '
        'C\'est le rendu typique des jeux de stratégie et du pixel-art.', 'key')
    body(doc, 'Chaque shape est un ensemble de voxels — des cubes élémentaires placés sur une grille 3D. '
         'Dans la projection isométrique, un voxel devient simplement un petit carré (fillRect) dessiné '
         'à une position 2D calculée par la formule isométrique.')
    body(doc, 'Le script crée un <canvas> de 120×120px par carte, l\'injecte dynamiquement dans le DOM, '
         'et gère l\'animation via requestAnimationFrame. '
         'L\'animation ne tourne que pendant le hover (avec fade-in/out fluide).')

    # ── S2 ──
    h2(doc, '02', 'La projection isométrique')
    body(doc, 'La projection isométrique est une vue axonométrique standard où l\'axe Y va "vers le haut" sur l\'écran, '
         'et les axes X et Z partent en diagonale à 30° sur la gauche et la droite. '
         'Ici, une rotation Y dynamique est ajoutée pour faire tourner la shape lors du hover.')

    h3(doc, 'Formule complète annotée')
    code(doc, '''\
function project(x, y, z, angle, scale) {

  // ── ÉTAPE 1 : Rotation autour de l\'axe Y ─────────────────────────
  // On fait pivoter le point dans le plan horizontal (X-Z)
  // avant de projeter. C\'est ce qui crée la rotation visible.
  const rx =  x * Math.cos(angle) + z * Math.sin(angle);  // X après rotation
  const rz = -x * Math.sin(angle) + z * Math.cos(angle);  // Z après rotation
  // Note : Y ne change pas (on ne tourne que autour de l\'axe vertical)

  // ── ÉTAPE 2 : Projection isométrique ─────────────────────────────
  // Formule isométrique standard :
  //   sx = (rx - rz) × scale × cos(30°)
  //   sy = (rx + rz) × scale × sin(30°) - y × scale
  // cos(30°) ≈ 0.866  |  sin(30°) = 0.5
  const sx = (rx - rz) * scale * 0.866;
  const sy = (rx + rz) * scale * 0.5 - y * scale;

  // ── ÉTAPE 3 : Profondeur pour le tri ─────────────────────────────
  // Sert uniquement au painter\'s algorithm.
  const depth = rx + rz - y;

  return { sx, sy, depth };
}''')

    h3(doc, 'Différence avec la perspective classique (hero-canvas)')
    table_2col(doc,
        ['', 'Perspective (hero-canvas)', 'Isométrique (hiw-voxel)'],
        [
            ['Taille selon Z',       'Diminue avec la distance',          'Constante (pas de division par Z)'],
            ['Lignes parallèles',    'Convergent vers un point de fuite', 'Restent parallèles'],
            ['Point de dessin',      'ctx.arc() (cercle)',                'ctx.fillRect() (carré)'],
            ['Style visuel',         'Réaliste, cinématique',             'Pixel-art, jeu de stratégie'],
            ['Rotation',             'Y + X (deux axes)',                 'Y uniquement'],
        ]
    )
    box(doc, 'Pourquoi 0.866 et 0.5 ? Ce sont cos(30°) et sin(30°). Dans une vue isométrique standard à 30°, '
        'les axes X et Z sont inclinés à 30° par rapport à l\'horizontale. Le facteur 0.866 "aplatit" horizontalement '
        'et 0.5 "comprime" verticalement. Ensemble, ils créent l\'angle caractéristique de l\'iso.', 'tip')

    # ── S3 ──
    h2(doc, '03', 'Les 4 shapes — Détail de chaque animation')
    table_2col(doc,
        ['Carte', 'Nom', 'Description', 'Couleur'],
        [
            ['01 — Choose Plan',      'Cube creux',       'Représente un produit/package.',      '#10B981 (émeraude)'],
            ['02 — Pass Challenge',   'Escalier 4 marches','Représente la progression.',         '#34d399 (émeraude clair)'],
            ['03 — Verify Identity',  'Grille de scan',   'Représente un document à vérifier.', '#4a9eff (bleu)'],
            ['04 — Get Funded',       'Bar chart',        'Représente la croissance.',           '#f59e0b (ambre)'],
        ]
    )

    h3(doc, 'Carte 01 — Cube creux (buildCube)')
    body(doc, 'Le cube couvre une grille 5×5×5 (N=4, de 0 à 4). Un voxel appartient au cube uniquement s\'il est sur '
         'au moins 2 faces simultanément — ce qui sélectionne les arêtes et les coins, donnant un effet wireframe.')
    code(doc, '''\
function buildCube() {
  const pts = [];
  const N = 4;   // taille du cube (4 unités de côté)

  for (let x = 0; x <= N; x++) {
    for (let y = 0; y <= N; y++) {
      for (let z = 0; z <= N; z++) {

        // Compter sur combien de faces ce voxel se trouve
        // (x=0 = face gauche, x=N = face droite, etc.)
        const onFaces = [x===0, x===N, y===0, y===N, z===0, z===N]
          .filter(Boolean).length;

        // onFaces === 1 → centre d\'une face (exclu)
        // onFaces === 2 → sur une arête (inclus)
        // onFaces === 3 → coin du cube (inclus)
        if (onFaces >= 2) pts.push([x - N/2, y - N/2, z - N/2]);
        // Soustraction de N/2 pour centrer le cube à l\'origine
      }
    }
  }
  return pts;
}''')
    box(doc, 'Résultat : 98 points placés sur les 12 arêtes et 8 coins du cube, centré en (0,0,0). '
        'Les faces intérieures sont vides — d\'où l\'effet "cage de fil".', 'info')

    h3(doc, 'Carte 02 — Escalier (buildStairs)')
    body(doc, '4 marches ascendantes. Chaque marche S a une surface supérieure (une ligne de points Y=S) '
         'et deux faces latérales (avant et arrière, le long de Z). L\'escalier monte en diagonale X=Y.')
    code(doc, '''\
function buildStairs() {
  const pts = [];
  const S = 4;   // nombre de marches

  for (let s = 0; s < S; s++) {

    // Surface supérieure de la marche s
    // X = s (avance d\'une unité par marche)
    // Y = s (monte d\'une unité par marche)
    // Z parcourt toute la profondeur
    for (let z = 0; z < S; z++) {
      pts.push([s, s, z]);
    }

    // Faces latérales (avant Z=0 et arrière Z=S-1)
    for (let y = 0; y <= s; y++) {
      pts.push([s, y, 0]);       // face avant
      pts.push([s, y, S - 1]);   // face arrière
    }
  }

  // Centrer l\'escalier à l\'origine (0,0,0)
  // Y est inversé (le négatif va "vers le haut" visuellement)
  return pts.map(([x, y, z]) => [
    x - (S-1)/2,
    -(y - (S-1)/2),   // inversion Y → la marche la plus haute est en haut
    z - (S-1)/2
  ]);
}''')
    box(doc, 'Pourquoi l\'inversion Y ? Dans le système de coordonnées 3D utilisé ici, Y positif pointe vers le bas à l\'écran. '
        'Sans l\'inversion, l\'escalier monterait visuellement vers le bas. '
        'Le -(y - ...) corrige ce comportement pour que les marches montent vers le haut.', 'warn')

    h3(doc, 'Carte 03 — Grille de scan (buildScanGrid)')
    body(doc, 'Une grille plate 6×6 représentant un document, avec deux lignes surélevées au centre simulant le faisceau d\'un scanner.')
    code(doc, '''\
function buildScanGrid() {
  const pts = [];

  // ── Grille plate 6×6 à Y=0 (le "document") ───────────────────────
  for (let x = 0; x <= 5; x++) {
    for (let z = 0; z <= 5; z++) {
      pts.push([x, 0, z]);
    }
  }

  // ── Lignes de scan surélevées (Y=1 et Y=2) ───────────────────────
  // Seulement au centre de la grille (X de 1 à 4, Z de 2 à 3)
  for (let x = 1; x <= 4; x++) {
    pts.push([x, 1, 2]);   // 1ère ligne surélevée, rang avant
    pts.push([x, 1, 3]);   // 1ère ligne surélevée, rang arrière
    pts.push([x, 2, 2]);   // 2ème ligne surélevée, rang avant
    pts.push([x, 2, 3]);   // 2ème ligne surélevée, rang arrière
  }

  // Centrer à l\'origine (−2.5 sur X et Z)
  return pts.map(([x, y, z]) => [x - 2.5, y, z - 2.5]);
}''')

    h3(doc, 'Carte 04 — Bar chart (buildBarChart)')
    body(doc, 'Cinq barres de hauteurs croissantes (1, 2, 3, 4, 5 unités) représentant une courbe de croissance. '
         'Chaque barre est un prisme 1×H×1 dont les 4 colonnes de coins sont tracées.')
    code(doc, '''\
function buildBarChart() {
  const heights = [1, 2, 3, 4, 5];   // hauteur de chaque barre
  const pts = [];
  const maxH = Math.max(...heights);  // = 5, pour centrer verticalement

  heights.forEach((h, i) => {
    const bx = i * 1.4;   // espacement horizontal : 1.4 unité entre barres

    // ── Colonnes verticales de la barre (4 coins) ──────────────────
    for (let y = 0; y < h; y++) {
      pts.push([bx,     y, 0]);   // coin avant gauche
      pts.push([bx,     y, 1]);   // coin arrière gauche
      pts.push([bx + 1, y, 0]);   // coin avant droit
      pts.push([bx + 1, y, 1]);   // coin arrière droit
    }
    // ── Chapeau de la barre (surface supérieure, 4 points) ─────────
    pts.push([bx,     h, 0]);
    pts.push([bx,     h, 1]);
    pts.push([bx + 1, h, 0]);
    pts.push([bx + 1, h, 1]);
  });

  // Centrer l\'ensemble
  const totalW = (heights.length - 1) * 1.4 + 1;
  return pts.map(([x, y, z]) => [
    x - totalW / 2,     // centrer horizontalement
    y - maxH / 2,       // centrer verticalement
    z - 0.5             // centrer en profondeur
  ]);
}''')

    # ── S4 — RENDERER ──
    h2(doc, '04', 'Le renderer — Boucle d\'animation et états')
    body(doc, 'Chaque canvas a son propre renderer créé par createRenderer(canvas, cfg). '
         'Le renderer maintient 4 variables d\'état internes qui pilotent l\'animation :')
    code(doc, '''\
let angle         = cfg.startAngle;  // angle de rotation initial (rad)
let opacity       = 0;               // opacité courante (0 = invisible)
let targetOpacity = 0;               // opacité cible (0 ou 0.55)
let hovered       = false;           // la carte est-elle survolée ?
let raf           = null;            // handle requestAnimationFrame''')

    h3(doc, 'Fade-in / Fade-out')
    body(doc, 'La transition d\'opacité est une interpolation exponentielle — l\'opacité courante s\'approche de la cible '
         'de 9% par frame. C\'est doux mais plus rapide au début, plus lent à la fin (ease-out naturel).')
    code(doc, '''\
// Dans tick(), chaque frame :
const dOp = (targetOpacity - opacity) * 0.09;  // 9% de la distance restante

if (Math.abs(dOp) > 0.002) {
  opacity += dOp;   // approche progressive de la cible
  dirty = true;
} else if (targetOpacity === 0 && opacity < 0.01) {
  opacity = 0;      // snap à 0 pour éviter les valeurs résiduelles
}

// Valeurs de targetOpacity :
// Hover entrant  → targetOpacity = 0.55  (semi-transparent)
// Hover sortant  → targetOpacity = 0     (disparaît)''')

    h3(doc, 'Rotation au hover')
    body(doc, 'La rotation ne se produit que pendant le hover. Quand la souris quitte la carte, la rotation s\'arrête '
         'mais la shape reste à son angle courant (pas de retour à l\'angle initial).')
    code(doc, '''\
if (hovered) {
  angle += 0.014;   // ~0.8°/frame, ~48°/seconde à 60fps
  dirty = true;     // force un redraw
}
// Sans hover : angle ne change pas, rotation figée''')

    h3(doc, 'Glow effect — Lueur couleur')
    body(doc, 'L\'effet de lueur est obtenu avec ctx.shadowBlur. Il s\'applique à tous les fillRect de la frame.')
    code(doc, '''\
ctx.shadowColor = cfg.color;   // ex : \'#10B981\' pour émeraude
ctx.shadowBlur  = 10;          // rayon de la lueur en pixels
ctx.fillStyle   = cfg.color;

pts.forEach(({ sx, sy }) => {
  ctx.fillRect(cx + sx - ds/2, cy + sy - ds/2, ds, ds);
  // ds = dotSize = 2.8px → chaque voxel est un carré 2.8×2.8
});

ctx.shadowBlur  = 0;   // réinitialiser après le dessin''')
    box(doc, 'Performance : shadowBlur est coûteux en CPU. Il est désactivé (= 0) immédiatement après le dessin. '
        'De plus, le renderer suspend complètement la boucle RAF quand la shape est invisible (opacity < 0.005), '
        'ce qui économise des ressources quand aucune carte n\'est survolée.', 'warn')

    h3(doc, 'Optimisation : pause quand inactif')
    code(doc, '''\
function tick() {
  let dirty = false;

  if (hovered) { angle += 0.014; dirty = true; }

  const dOp = (targetOpacity - opacity) * 0.09;
  if (Math.abs(dOp) > 0.002) { opacity += dOp; dirty = true; }

  if (dirty || opacity > 0.005) {
    draw();
    raf = requestAnimationFrame(tick);   // continuer la boucle
  } else {
    raf = null;   // ← PAUSE : plus de RAF, zéro CPU utilisé
  }
}

// Pour redémarrer après une pause :
function ensure() {
  if (!raf) raf = requestAnimationFrame(tick);  // relance seulement si stoppé
}''')

    # ── S5 ──
    h2(doc, '05', 'Le Painter\'s Algorithm — Tri par profondeur')
    body(doc, 'Pour que les voxels proches soient dessinés par-dessus les voxels lointains (illusion de profondeur), '
         'tous les points sont triés par leur valeur de depth avant d\'être dessinés.')
    code(doc, '''\
// 1. Projeter tous les points 3D → 2D + calculer leur profondeur
const pts = cfg.shape.map(([x, y, z]) => project(x, y, z, angle, cfg.scale));

// 2. Trier par profondeur croissante (les plus lointains en premier)
pts.sort((a, b) => a.depth - b.depth);

// 3. Dessiner dans cet ordre : les lointains d\'abord, proches par-dessus
pts.forEach(({ sx, sy }) => {
  ctx.fillRect(cx + sx - ds/2, cy + sy - ds/2, ds, ds);
});''')
    body(doc, 'La valeur de profondeur utilisée est depth = rx + rz - y. Cette formule projette les coordonnées 3D '
         'sur un axe diagonal isométrique, donnant une mesure de "distance à l\'œil" cohérente avec la vue iso.')

    # ── S6 ──
    h2(doc, '06', 'Intégration HTML/CSS')
    h3(doc, 'Structure HTML requise')
    body(doc, 'Le script cherche automatiquement tous les éléments .hiw-card dans la page et leur injecte un canvas '
         'en coin supérieur droit. Les 4 premières cartes trouvées reçoivent respectivement le cube, l\'escalier, la grille, et le bar chart.')
    code(doc, '''\
<!-- La structure minimale pour une carte -->
<div class="hiw-card">
  <div class="hiw-num">01</div>
  <h3>Choose Plan</h3>
  <p>Description...</p>
  <!-- Le canvas .hiw-voxel est injecté ICI automatiquement par le JS -->
</div>''')
    h3(doc, 'CSS du canvas (dans main.css)')
    code(doc, '''\
.hiw-voxel {
  position: absolute;
  top: 12px;
  right: 12px;
  pointer-events: none;   /* ne bloque pas les clics */
  z-index: 2;
  /* Le JS fixe width/height à 120×120px via style inline */
}''')
    box(doc, 'Important : La carte parente (.hiw-card) doit avoir position: relative (ou autre valeur non-static) '
        'pour que le position: absolute du canvas se positionne correctement en son sein.', 'info')
    h3(doc, 'Chargement du script')
    code(doc, '''\
<!-- Dans footer.php, chargé uniquement sur la page home -->
<?php if ($currentPage === \'home\'): ?>
<script defer src="assets/js/hiw-voxel.js"></script>
<?php endif; ?>''')

    # ── S7 ──
    h2(doc, '07', 'Configuration des 4 cartes (CONFIGS)')
    code(doc, '''\
const CONFIGS = [
  { shape: buildCube(),      color: \'#10B981\', scale: 13, dotSize: 2.8, startAngle: 0.5  },
  { shape: buildStairs(),    color: \'#34d399\', scale: 13, dotSize: 2.8, startAngle: 0.3  },
  { shape: buildScanGrid(),  color: \'#4a9eff\', scale: 11, dotSize: 2.8, startAngle: 0.6  },
  { shape: buildBarChart(),  color: \'#f59e0b\', scale: 11, dotSize: 2.8, startAngle: 0.4  },
];''')
    table_2col(doc,
        ['Paramètre', 'Type', 'Rôle'],
        [
            ['shape',      'Array',  'Tableau de points [x,y,z] retourné par la fonction de build'],
            ['color',      'String', 'Couleur CSS des voxels ET de la lueur (shadowColor)'],
            ['scale',      'Number', 'Facteur d\'échelle isométrique. 13 = grande shape, 11 = légèrement plus petite'],
            ['dotSize',    'Number', 'Taille du carré d\'un voxel en pixels (2.8px)'],
            ['startAngle', 'Number', 'Angle de rotation initial en radians (chaque carte commence différemment)'],
        ]
    )
    body(doc, 'Tous les canvas sont fixes à 120×120px (logiques) avec support du devicePixelRatio (max 2×) '
         'pour les écrans Retina. En pratique le canvas physique peut faire 240×240px sur un écran Retina.')

    # ── S8 ──
    h2(doc, '08', 'Comportement mobile')
    body(doc, 'Sur les appareils tactiles, il n\'y a pas d\'événement mouseenter. Le script détecte automatiquement le mobile :')
    code(doc, '''\
const isMobile = window.matchMedia(\'(hover: none) and (pointer: coarse)\').matches;
// hover: none  → l\'appareil n\'a pas de survol (pas de souris)
// pointer: coarse → pointeur imprécis (doigt, pas souris)''')
    body(doc, 'Sur mobile, un IntersectionObserver remplace le hover : quand une carte entre dans le viewport '
         '(à 50% de visibilité), l\'animation s\'active. Les autres cartes se désactivent, '
         'créant un effet "une seule active à la fois".')
    code(doc, '''\
const io = new IntersectionObserver((entries) => {
  entries.forEach(entry => {
    const match = renderers.find(r => r.card === entry.target);
    if (!match) return;

    if (entry.isIntersecting) {
      renderers.forEach(r => r.renderer.leave());  // désactiver toutes les autres
      match.renderer.enter();                       // activer celle-ci
    } else {
      match.renderer.leave();                       // désactiver en quittant le viewport
    }
  });
}, { threshold: 0.5 });   // déclencher à 50% de visibilité

renderers.forEach(r => io.observe(r.card));''')

    # ── S9 ──
    h2(doc, '09', 'Checklist — Créer une nouvelle shape from scratch')
    body(doc, 'Pour ajouter une 5ème carte avec une nouvelle animation (ex: une flèche, une étoile, une maison) :')
    step(doc, 1, 'Écrire la fonction buildXxx()',
         'Retourner un tableau de [x, y, z] en coordonnées centrées à l\'origine. Y positif = bas, Y négatif = haut à l\'écran.')
    step(doc, 2, 'Tester la forme avec angle=0',
         'Visualiser d\'abord sans rotation pour vérifier que la shape est centrée et lisible de face.')
    step(doc, 3, 'Ajuster scale et dotSize',
         'Si la shape déborde du canvas 120×120px, diminuer scale (13 → 10). Si les points sont trop petits, augmenter dotSize.')
    step(doc, 4, 'Choisir un startAngle',
         'Valeur entre 0 et 2π. Choisir un angle qui présente la face la plus "lisible" de la shape au premier coup d\'œil.')
    step(doc, 5, 'Ajouter l\'entrée dans CONFIGS[]',
         '{ shape: buildXxx(), color: \'#hexcode\', scale: 12, dotSize: 2.8, startAngle: 0.4 }')
    step(doc, 6, 'Ajouter la 5ème carte en HTML',
         'Ajouter un <div class="hiw-card"> à la suite des 4 existants. Le JS l\'associe automatiquement à CONFIGS[4].')
    step(doc, 7, 'Vérifier le grid CSS',
         'Le grid est actuellement grid-template-columns: repeat(4, 1fr) — changer à repeat(5, 1fr) pour accommoder 5 colonnes.')
    box(doc, 'Astuce pour déboguer : Commencer par un simple carré 2D ([[0,0,0],[1,0,0],[1,0,1],[0,0,1]]) pour vérifier '
        'que la projection s\'affiche correctement avant de construire une shape complexe.', 'tip')
    box(doc, 'Récapitulatif des différences :\n'
        'hero-canvas = perspective + cercles (arc()) + physique spring + dot grid + shapes ~900 pts\n'
        'hiw-voxel   = isométrique + carrés (fillRect()) + fade + hover + pas de dot grid + shapes ~50–100 pts', 'info')

    footer(doc)
    path = os.path.join(BASE, 'HIW_Voxel_Animations_Isometriques.docx')
    doc.save(path)
    print(f'Saved: {path}')


if __name__ == '__main__':
    create_hero_canvas_doc()
    create_hiw_voxel_doc()
    print('Done.')
